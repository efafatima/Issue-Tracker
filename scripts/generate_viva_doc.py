from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


OUTPUT = "BZU_Issue_Tracker_External_Viva_Guide.docx"


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = 1
        r = p.add_run(subtitle)
        r.font.size = Pt(12)


def h(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text=""):
    doc.add_paragraph(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def nums(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        t.rows[0].cells[i].text = header
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def section_intro(doc, english, hinglish):
    p(doc, english)
    p(doc, "Hinglish: " + hinglish)


def build_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    add_title(
        doc,
        "BZU Issue Tracker / Complaint Management System",
        "Final Year Project External Viva Preparation Guide - English + Roman Urdu/Hinglish",
    )
    p(doc, "Prepared for BS Computer Science external viva. This document explains the project like an examiner expects: problem, architecture, database, security, workflow, code logic, demo, presentation, and viva questions.")
    p(doc, "Important note: The codebase currently defines these roles in the database: Student, Faculty Member, HOD, DSA, Supervisor. If you mention Admin in viva, explain it as a general administrative category, while actual implemented roles are HOD, DSA, Faculty Member, and Supervisor.")

    h(doc, "Quick 60-Second Project Introduction", 1)
    p(doc, "My project is BZU Issue Tracker, a web-based Complaint Management System for Bahauddin Zakariya University. It allows students to submit complaints, track progress, receive notifications, and give feedback. Staff roles such as HOD, DSA, Faculty Member, and Supervisor review complaints, assign them to faculty, update status, resolve issues, and monitor analytics. The system is built with Next.js App Router on the frontend/backend layer and Supabase for PostgreSQL database, authentication, storage, and secure data access.")
    p(doc, "Hinglish: Mera project BZU ke liye complaint management system hai. Students complaints submit aur track karte hain. HOD/DSA complaints review karte hain, faculty ko assign karte hain, faculty resolve karti hai, aur supervisor analytics aur department performance monitor karta hai.")

    h(doc, "Table of Contents", 1)
    nums(doc, [
        "Project Overview",
        "System Architecture",
        "Folder Structure",
        "Database Design",
        "Authentication and Authorization",
        "Complete Complaint Flow",
        "Code Logic",
        "Next.js Concepts",
        "React Concepts",
        "Supabase Concepts",
        "Security",
        "Performance",
        "Reports and Analytics",
        "Presentation Plan",
        "Demo Script",
        "Tricky External Questions",
        "100+ Viva Questions",
        "Final Revision Checklist",
    ])

    h(doc, "Part 1 - Project Overview", 1)
    section_intro(
        doc,
        "The project solves the problem of manual, slow, and untraceable complaint handling in a university environment.",
        "Ye system manual complaint process ko digital, trackable aur role-based banata hai.",
    )
    h(doc, "Problem Statement", 2)
    p(doc, "In many universities, students submit complaints verbally, through paper forms, or through informal communication. This creates delays, missing records, unclear responsibility, and no transparent tracking for students.")
    h(doc, "Why This Project Was Needed", 2)
    bullets(doc, [
        "Students need a proper channel to report academic, administrative, facilities, behavior-related, and other issues.",
        "University staff need a structured workflow for review, assignment, resolution, and reporting.",
        "Management needs analytics to identify repeated problems and department performance.",
        "A digital record improves accountability and reduces communication gaps.",
    ])
    h(doc, "Existing Problems", 2)
    bullets(doc, [
        "No central complaint history.",
        "Students do not know current complaint status.",
        "Complaints can be delayed because ownership is unclear.",
        "Paper/manual reports are hard to analyze.",
        "No automatic notifications or activity log.",
    ])
    h(doc, "Proposed Solution", 2)
    p(doc, "A role-based web application where students submit complaints and the system routes them to the relevant authority. HOD/DSA review complaints, assign accepted complaints to faculty, faculty resolves them, and the supervisor monitors overall progress through dashboards and analytics.")
    h(doc, "Objectives", 2)
    bullets(doc, [
        "Provide secure registration and login.",
        "Allow students to submit and track complaints.",
        "Route complaints by category to HOD or DSA.",
        "Support assignment to faculty members.",
        "Maintain notifications and activity logs.",
        "Provide dashboards, weekly reports, and department-wise analytics.",
        "Improve transparency, accountability, and response time.",
    ])
    h(doc, "Scope", 2)
    p(doc, "The system covers complaint submission, review, assignment, resolution, feedback, notifications, attachment support, analytics, department management, and role-based access control.")
    h(doc, "Future Enhancements", 2)
    bullets(doc, [
        "Mobile app for students and staff.",
        "Real-time notifications using Supabase Realtime.",
        "Advanced AI classification using NLP models.",
        "SLA-based escalation rules per department.",
        "SMS/WhatsApp integration.",
        "Exportable PDF reports.",
        "Public complaint trend dashboard without exposing personal data.",
    ])

    h(doc, "Part 2 - System Architecture", 1)
    section_intro(
        doc,
        "The architecture is a modern full-stack Next.js application using Supabase as backend service.",
        "Next.js user interface aur API routes handle karta hai, Supabase auth/database/storage handle karta hai.",
    )
    code(doc, """
Student / Staff Browser
        |
        v
Next.js App Router UI
        |
        v
Next.js API Routes (/api/...)
        |
        v
Supabase Auth + PostgreSQL + Storage
        |
        v
Response returned to dashboard
""")
    h(doc, "Client", 2)
    p(doc, "The client is the browser UI built with React components inside Next.js. It renders login, registration, dashboard, complaint forms, cards, charts, profile editor, and role-specific views.")
    h(doc, "Server", 2)
    p(doc, "The server layer is implemented with Next.js API routes under src/app/api. These routes verify the logged-in user, check role permissions, query Supabase, and return JSON responses.")
    h(doc, "Supabase", 2)
    p(doc, "Supabase provides PostgreSQL database, Supabase Auth, file storage for complaint attachments, Row Level Security policies, and service-role access for secure server-side workflows.")
    h(doc, "Authentication Flow", 2)
    code(doc, """
User enters email/password
        |
Supabase Auth verifies credentials
        |
Session/JWT is created
        |
Frontend sends Bearer token to API route
        |
API route calls supabase.auth.getUser(token)
        |
Profile is loaded from profiles table
        |
Role-based decision is applied
""")
    h(doc, "Request/Response Flow", 2)
    p(doc, "Example: when a student submits a complaint, the dashboard sends a POST request to /api/complaints. The API validates the user, validates input, predicts category/priority if needed, inserts the complaint, creates activity log and notification, then returns the created complaint.")

    h(doc, "Part 3 - Folder Structure", 1)
    table(doc, ["Folder/File", "Purpose"], [
        ("src/app", "Next.js App Router pages, layouts, and API route handlers."),
        ("src/app/page.jsx", "Welcome/landing screen with Login and Register links."),
        ("src/app/login/page.jsx", "Login page for authenticated access."),
        ("src/app/register/page.jsx", "Registration page for new users."),
        ("src/app/dashboard/page.jsx", "Main role-based dashboard for students, faculty, HOD, DSA, and supervisor."),
        ("src/app/dashboard/analytics/page.jsx", "Route entry for analytics view."),
        ("src/app/dashboard/complaints/page.jsx", "Route entry for complaint list view."),
        ("src/app/dashboard/departments/page.jsx", "Supervisor department insights route."),
        ("src/app/api", "Backend API routes for auth, complaints, users, departments, bot, reports, analytics."),
        ("src/components", "Reusable UI components such as ComplaintForm, ComplaintCard, StatusBadge, StatCard, DepartmentManager."),
        ("src/lib", "Shared business logic: auth, API helpers, Supabase client, workflow rules, ML helper, email, attachments, escalation."),
        ("supabase/schema.sql", "Database schema, enums, tables, indexes, RLS policies, storage bucket."),
        ("public", "Static assets such as BZU logo."),
        ("package.json", "Project dependencies and scripts."),
    ])
    p(doc, "Hinglish: app folder pages aur APIs ke liye hai, components reusable UI ke liye, lib shared logic ke liye, aur supabase folder database setup ke liye.")

    h(doc, "Part 4 - Database Design", 1)
    section_intro(
        doc,
        "The database is relational because complaints, users, departments, assignments, comments, notifications, and logs are connected entities.",
        "Relational database is liye use hui kyun ke data relationships clear hain: user ki complaints, department, assigned teacher, comments, notifications.",
    )
    h(doc, "Enums", 2)
    table(doc, ["Enum", "Values", "Why"], [
        ("user_role", "Student, Faculty Member, HOD, DSA, Supervisor", "Controls role-based access."),
        ("complaint_status", "Submitted, In Progress, Resolved, Closed, Rejected, Escalated", "Represents complaint lifecycle."),
        ("complaint_category", "Academic, Administrative, Facilities, Behavior-related, Other", "Routes complaints to correct authority."),
        ("complaint_priority", "Low, Medium, High, Urgent", "Shows urgency."),
        ("routed_role", "HOD, DSA, Supervisor", "Defines which role should review complaint."),
    ])
    h(doc, "Tables", 2)
    table(doc, ["Table", "Main Columns", "Purpose"], [
        ("profiles", "id, username, email, role, department_id, faculty_designation, is_active", "Stores application profile linked with Supabase auth.users."),
        ("departments", "id, name, hod_id, dsa_id", "Stores departments and their responsible HOD/DSA."),
        ("complaints", "id, user_id, title, description, category, priority, department_id, routed_to_role, status, assigned_teacher_id, rating", "Main complaint records."),
        ("complaint_attachments", "id, complaint_id, file_path, file_url, file_type", "Stores uploaded evidence metadata."),
        ("complaint_comments", "id, complaint_id, user_id, description, is_internal", "Stores comments and internal notes."),
        ("notifications", "id, user_id, complaint_id, message, is_read, is_sent", "Stores user notifications."),
        ("category_routes", "id, category, default_role, department_id", "Maps complaint category to reviewing role."),
        ("activity_logs", "id, complaint_id, user_id, action, old_value, new_value, ip_address", "Audit trail for accountability."),
    ])
    h(doc, "Important Relationships", 2)
    bullets(doc, [
        "profiles.id references auth.users(id): every app profile belongs to an authenticated Supabase user.",
        "profiles.department_id references departments.id: user belongs to a department.",
        "departments.hod_id and departments.dsa_id reference profiles.id: department has responsible staff.",
        "complaints.user_id references profiles.id: complaint belongs to student.",
        "complaints.assigned_teacher_id references profiles.id: complaint can be assigned to faculty.",
        "complaint_attachments, complaint_comments, notifications, activity_logs reference complaints.id.",
    ])
    h(doc, "Indexes and Constraints", 2)
    bullets(doc, [
        "Unique username/email avoids duplicate profiles.",
        "Rating check ensures rating is between 1 and 5.",
        "Indexes on role/department, complaint user, status/priority, assignee, notifications, and activity logs improve query speed.",
        "RLS is enabled on all important tables. Server-side APIs use service-role key for controlled workflows.",
    ])

    h(doc, "Part 5 - Authentication and Authorization", 1)
    p(doc, "Authentication means proving who the user is. Authorization means checking what that user is allowed to do.")
    p(doc, "In this project, Supabase Auth manages login sessions. The API reads the Bearer token, verifies it using supabase.auth.getUser(token), then loads the user's profile and role from the profiles table.")
    h(doc, "Student Login", 2)
    p(doc, "Student logs in through the UI. Supabase returns a session. The dashboard checks session and loads profile. Student can create complaints and view only own complaints.")
    h(doc, "Staff Login", 2)
    p(doc, "HOD, DSA, Faculty Member, and Supervisor also authenticate through Supabase. Their permissions come from profiles.role.")
    h(doc, "RBAC Rules", 2)
    table(doc, ["Role", "Allowed Actions"], [
        ("Student", "Submit complaints, view own complaints, track status, rate/feedback after resolution."),
        ("HOD", "Review Academic and Behavior-related complaints for own department, accept/reject, assign faculty."),
        ("DSA", "Review Administrative, Facilities, and Other complaints for own department, accept/reject, assign faculty, weekly reports."),
        ("Faculty Member", "View assigned complaints, update assigned complaint status to Resolved, add comments/evidence."),
        ("Supervisor", "View overall activity, department insights, analytics, assignment oversight."),
    ])
    h(doc, "How Unauthorized Users Are Blocked", 2)
    bullets(doc, [
        "API routes call currentUser(request). If no token exists, response is 401 Authentication required.",
        "If token is invalid, response is 401 Invalid session.",
        "If profile is missing, response is 403 Profile not found.",
        "Specific endpoints check role. For example, only Student can POST /api/complaints.",
        "Workflow helpers such as canViewComplaint and scopedComplaintQuery restrict records by role.",
    ])

    h(doc, "Part 6 - Complete Complaint Flow", 1)
    code(doc, """
Student Login
   -> Complaint Form
   -> Client Validation
   -> POST /api/complaints
   -> Server Authentication
   -> Category/Priority Suggestion
   -> Insert into complaints
   -> Insert activity_logs
   -> Insert notifications
   -> Send emails
   -> Dashboard update
   -> HOD/DSA Review
   -> Assign Faculty
   -> Faculty marks Resolved
   -> Final review / reports / analytics
""")
    nums(doc, [
        "Student signs in and dashboard loads the profile.",
        "Student fills title, description, category, priority, department, anonymity option, and attachments if used.",
        "Frontend validates required fields.",
        "POST /api/complaints authenticates the Bearer token.",
        "The route checks role equals Student.",
        "Title and description are trimmed and validated.",
        "Category and priority are selected or predicted using keyword-based helper functions.",
        "Complaint is inserted with status Submitted.",
        "Activity log records submission and routing.",
        "Notification is created for the student.",
        "Emails may be sent to relevant users.",
        "HOD/DSA reviews complaint using /review route.",
        "Accepted complaint can be assigned to a Faculty Member using /assign route.",
        "Assignment changes status to In Progress and notifies teacher/student.",
        "Assigned faculty marks complaint Resolved using /status route.",
        "Reports and analytics summarize status counts, weekly activity, and department history.",
    ])

    h(doc, "Part 7 - Code Logic", 1)
    table(doc, ["File", "Why It Exists"], [
        ("src/lib/auth.js", "Centralizes token verification and profile loading."),
        ("src/lib/workflow.js", "Stores role-based complaint visibility and query scoping rules."),
        ("src/lib/ml.js", "Provides simple keyword-based category, severity, and similarity suggestions."),
        ("src/lib/api.js", "Common API response and JSON parsing helpers."),
        ("src/lib/clientApi.js", "Frontend helper for calling API routes with session token."),
        ("src/lib/supabaseClient.js", "Creates Supabase clients for browser/server usage."),
        ("src/lib/attachments.js", "Handles signed URLs and attachment access."),
        ("src/lib/escalation.js", "Escalates overdue complaints based on deadlines."),
        ("src/components/ComplaintForm.jsx", "Collects complaint data from student."),
        ("src/components/ComplaintCard.jsx", "Displays complaint details and role-based actions."),
        ("src/components/DepartmentManager.jsx", "Supervisor tools and department management."),
        ("src/components/StatusBadge.jsx", "Consistent status UI."),
        ("src/components/StatCard.jsx", "Reusable dashboard statistics card."),
        ("src/app/dashboard/page.jsx", "Main dashboard controller with role-specific UI."),
    ])
    p(doc, "Why components are separated: It improves readability, reuse, testing, and maintenance. For example, ComplaintCard can be reused in dashboard and complaint list without duplicating UI.")
    p(doc, "Why state is used: Dashboard state stores profile, complaints, stats, analytics, notifications, selected complaint, loading and errors. React re-renders UI when this state changes.")
    p(doc, "Why dynamic imports are used: Heavy or client-only components such as chat/voice widgets and complaint forms are loaded only when needed, improving initial load.")

    h(doc, "Part 8 - Next.js Concepts in This Project", 1)
    table(doc, ["Concept", "Project Example"], [
        ("App Router", "Routes are inside src/app, e.g., /dashboard and /api/complaints."),
        ("Layout", "src/app/layout.jsx wraps pages and imports global CSS."),
        ("Page", "page.jsx files represent screens such as login, register, dashboard."),
        ("Client Component", "Dashboard uses 'use client' because it needs state, effects, router, and localStorage."),
        ("Server/API Route", "route.js files handle backend logic for complaints, users, auth, analytics."),
        ("Dynamic Route", "src/app/api/complaints/[id]/... handles complaint-specific actions."),
        ("Routing", "useRouter navigates between dashboard sections."),
        ("Metadata", "Can be used in layout/page for title and SEO if needed."),
        ("Loading/Error", "The dashboard currently uses custom loading/error UI state."),
    ])
    p(doc, "Exam answer for 'Why Next.js?': Next.js gives React UI plus backend API routes, file-based routing, server-side capabilities, optimized builds, and a clean full-stack structure. Plain React would need a separate backend server.")

    h(doc, "Part 9 - React Concepts in This Project", 1)
    table(doc, ["Concept", "Example"], [
        ("Components", "ComplaintForm, ComplaintCard, StatCard, StatusBadge."),
        ("Props", "ComplaintCard receives complaint, profile, teachers, onAction."),
        ("State", "Dashboard stores profile, complaints, stats, searchTerm, loading."),
        ("useEffect", "Loads session and profile when dashboard mounts."),
        ("useMemo", "Computes statCards and filtered complaints efficiently."),
        ("Controlled Forms", "Profile editor and faculty create form store input values in state."),
        ("Conditional Rendering", "Student sees complaint form; Faculty sees assigned complaints; Supervisor sees department manager."),
        ("Lists and Keys", "Complaints are mapped to ComplaintCard with key={complaint.id}."),
    ])

    h(doc, "Part 10 - Supabase Concepts", 1)
    bullets(doc, [
        "Database: PostgreSQL stores structured complaint data.",
        "Authentication: Supabase Auth manages users, sessions, and JWT tokens.",
        "Storage: complaint-attachments bucket stores uploaded files securely.",
        "RLS: Row Level Security is enabled to protect table-level access.",
        "Policies: Users can read own profile, students can read/create own complaints, users can read own notifications.",
        "Service role: Server-side API routes use service role for cross-role workflows while still applying custom permission checks.",
        "Indexes: Improve filtering by role, department, status, priority, assignee, and notification read status.",
    ])
    p(doc, "Exam answer for 'Why Supabase instead of Firebase?': Supabase gives PostgreSQL, SQL joins, relational constraints, RLS, and an open-source backend style. Firebase is strong for realtime NoSQL apps, but this project has relational data like users, departments, assignments, comments, and reports.")

    h(doc, "Part 11 - Security", 1)
    table(doc, ["Risk", "Protection in Project"], [
        ("SQL Injection", "Supabase query builder uses parameterized queries instead of manual SQL strings in API code."),
        ("Unauthenticated access", "API routes require Bearer token and verify it with Supabase Auth."),
        ("Unauthorized access", "Role checks and scopedComplaintQuery restrict actions and records."),
        ("Secrets exposure", "Service role keys must stay only in server environment variables, not frontend."),
        ("XSS", "React escapes rendered text by default; avoid dangerouslySetInnerHTML."),
        ("CSRF", "Bearer token API style reduces classic cookie-only CSRF risk, but secure session handling is still important."),
        ("File access", "Attachments are private and signed URLs are used for access."),
        ("Auditability", "activity_logs records important changes."),
    ])
    p(doc, "Important viva line: Authentication tells who the user is; authorization decides what the user can do.")

    h(doc, "Part 12 - Performance", 1)
    bullets(doc, [
        "Dynamic imports reduce initial JavaScript bundle for heavy/client-only components.",
        "Database indexes speed up common dashboard filters.",
        "scopedComplaintQuery limits returned data according to role.",
        "Charts and stats are computed from backend endpoints instead of scanning everything on the client.",
        "Future improvement: pagination for large complaint lists.",
        "Future improvement: caching department lists and analytics where data does not change every second.",
        "Future improvement: background jobs for escalation and email sending.",
    ])

    h(doc, "Part 13 - Reports and Analytics", 1)
    bullets(doc, [
        "Dashboard stat cards show totals and status-based counts.",
        "Analytics view shows complaint status trend.",
        "Notifications show recent complaint actions.",
        "Weekly report summarizes last 7 days: solved, pending, in progress, escalated, average rating, category counts.",
        "Department insights show department-level totals, in-process, completed, and pending complaints.",
        "Activity feed gives supervisor an audit trail.",
    ])

    h(doc, "Part 14 - 10 to 15 Minute Presentation Plan", 1)
    table(doc, ["Slide", "Title", "Bullet Points", "Speaker Notes", "Expected Question"], [
        (1, "Project Title", "BZU Issue Tracker; Complaint Management System; Next.js + Supabase", "Introduce yourself and project purpose.", "What problem does it solve?"),
        (2, "Problem Statement", "Manual complaints; no tracking; delayed response; weak reporting", "Explain pain points in university context.", "Why is manual process inefficient?"),
        (3, "Objectives", "Digital submission; RBAC; tracking; notifications; reports", "Show measurable goals.", "What are your main objectives?"),
        (4, "Users and Roles", "Student; HOD; DSA; Faculty; Supervisor", "Explain each role briefly.", "How is access controlled?"),
        (5, "System Architecture", "Browser; Next.js UI; API routes; Supabase Auth/Postgres/Storage", "Use architecture diagram.", "Why Next.js and Supabase?"),
        (6, "Database Design", "profiles; departments; complaints; notifications; logs", "Explain relational design.", "Why PostgreSQL?"),
        (7, "Complaint Workflow", "Submit; route; review; assign; resolve; report", "Walk through lifecycle.", "Who can update status?"),
        (8, "Security", "Auth; RBAC; RLS; env vars; audit logs", "Defend security choices.", "How unauthorized users are blocked?"),
        (9, "Analytics and Reports", "Status trends; weekly report; department report", "Explain management value.", "How are reports generated?"),
        (10, "Demo", "Student flow; admin flow; analytics", "Tell examiner what you will demonstrate.", "What if complaint is urgent?"),
        (11, "Limitations", "No mobile app; basic ML; limited realtime", "Be honest and professional.", "What would you improve?"),
        (12, "Future Work", "Mobile; realtime; advanced AI; SLA; exports", "End with growth roadmap.", "How will it scale?"),
    ])

    h(doc, "Part 15 - Demo Preparation Script", 1)
    table(doc, ["Step", "Action", "What To Say"], [
        (1, "Student Login", "I am logging in as a student to submit and track complaints securely."),
        (2, "Student Dashboard", "This dashboard shows student's own complaint stats and history."),
        (3, "Submit Complaint", "I enter title and description. The system validates input and can suggest category/priority."),
        (4, "Track Complaint", "After submission, the complaint appears with status Submitted and notification is created."),
        (5, "Logout", "Now I switch role to show administrative workflow."),
        (6, "HOD/DSA Login", "This role sees only complaints routed to their role and department."),
        (7, "Review Complaint", "The reviewer can accept or reject based on validity."),
        (8, "Assign Complaint", "Accepted complaint is assigned to an active faculty member."),
        (9, "Faculty Login", "Faculty sees assigned complaints only."),
        (10, "Update Status", "Faculty marks complaint as Resolved after handling it."),
        (11, "Notifications", "Student and staff receive notifications for important actions."),
        (12, "Reports", "Reports summarize weekly and department-wise complaint performance."),
        (13, "Analytics", "Analytics helps management identify trends and bottlenecks."),
    ])

    h(doc, "Part 16 - Tricky External Questions and Strong Answers", 1)
    table(doc, ["Question", "Strong Answer"], [
        ("Why Next.js instead of React?", "React is only UI. Next.js gives routing, API routes, server-side features, build optimization, and full-stack structure."),
        ("Why Supabase instead of Firebase?", "This project has relational data and reporting needs. Supabase provides PostgreSQL, SQL joins, constraints, RLS, and auth."),
        ("Why PostgreSQL?", "Complaints, users, departments, assignments, comments, and reports are relational, so PostgreSQL fits naturally."),
        ("Why RBAC?", "Different users have different responsibilities. RBAC prevents students from seeing staff data and prevents staff from unauthorized actions."),
        ("Why App Router?", "It is the modern Next.js routing model and organizes pages and API routes clearly."),
        ("Why not MongoDB?", "MongoDB is document-oriented. This project needs relationships, constraints, joins, and structured reports."),
        ("How to scale to 100,000 students?", "Add pagination, indexes, caching, background jobs, optimized queries, CDN, database monitoring, and possibly queue-based email/notification processing."),
        ("What if Supabase goes down?", "Show user-friendly error, retry, monitoring, backups, and disaster recovery plan. For production, use backups and high-availability options."),
        ("What is the weakest part?", "Basic keyword ML can be improved with trained NLP. Also realtime notifications and pagination can be enhanced."),
        ("How do you ensure data privacy?", "Auth, RBAC, RLS policies, private storage, signed URLs, and server-side checks."),
    ])

    h(doc, "Part 17 - 100+ Viva Questions", 1)
    easy = [
        "What is the name of your project?",
        "What problem does your project solve?",
        "Who are the main users?",
        "What is a complaint management system?",
        "What technologies did you use?",
        "What is Next.js?",
        "What is Supabase?",
        "What is PostgreSQL?",
        "What is authentication?",
        "What is authorization?",
        "What is RBAC?",
        "What roles exist in your project?",
        "What can a student do?",
        "What can HOD do?",
        "What can DSA do?",
        "What can Faculty Member do?",
        "What can Supervisor do?",
        "What is a dashboard?",
        "What is a notification?",
        "What is an activity log?",
        "What is a primary key?",
        "What is a foreign key?",
        "What is an API route?",
        "What is a React component?",
        "What is state in React?",
        "What is a prop in React?",
        "What is useEffect?",
        "What is useMemo?",
        "What is a controlled form?",
        "What is a database table?",
        "What is an enum?",
        "What is a status field?",
        "What is a complaint category?",
        "What is priority?",
        "Why do you need login?",
        "Why do you need department table?",
        "What is the purpose of profiles table?",
        "What is the purpose of complaints table?",
        "What is the purpose of notifications table?",
        "What is the purpose of attachments table?",
    ]
    medium = [
        "Explain complete complaint lifecycle.",
        "How does complaint routing work?",
        "How do you prevent a student from seeing other students' complaints?",
        "How does the API know which user is logged in?",
        "Why do you store profile separately from auth.users?",
        "How does Faculty Member update status?",
        "Why can only assigned faculty resolve a complaint?",
        "How do HOD and DSA differ?",
        "How is department-based access implemented?",
        "How are notifications generated?",
        "How are activity logs useful?",
        "How does your dashboard show role-specific UI?",
        "Why did you separate components?",
        "Why did you create lib/workflow.js?",
        "Why did you create lib/auth.js?",
        "How does scopedComplaintQuery work?",
        "How does canViewComplaint work?",
        "What happens when invalid token is sent?",
        "What happens when unauthorized role calls an endpoint?",
        "How are complaint attachments secured?",
        "How is input validation handled?",
        "What is Row Level Security?",
        "Why use service role on server side?",
        "What is the risk of exposing service role key?",
        "How do indexes improve performance?",
        "How are weekly reports useful?",
        "How is analytics useful for supervisor?",
        "How would you add pagination?",
        "How would you add realtime notifications?",
        "How do you handle errors in API routes?",
        "What is the difference between status and priority?",
        "Why store both suggested_category and category?",
        "Why store resolved_at?",
        "Why store assigned_by_id?",
        "Why use generated identity for complaint id?",
        "What is the role of storage bucket?",
        "How do you send emails?",
        "How do you escalate overdue complaints?",
        "What are the limitations of keyword-based ML?",
        "How would you test this project?",
    ]
    hard = [
        "Why is PostgreSQL better than MongoDB for this project?",
        "Explain normalization in your schema.",
        "What are possible race conditions in assignment?",
        "How would you avoid duplicate complaints?",
        "How would you implement SLA per category?",
        "How would you design audit logs for legal accountability?",
        "How would you secure file uploads from malware?",
        "How would you design backup and recovery?",
        "How would you scale to multiple campuses?",
        "How would you support 100,000 students?",
        "How would you optimize analytics queries?",
        "What data should be cached and what should not?",
        "How would you implement soft delete?",
        "How would you make the system multi-tenant?",
        "How would you handle Supabase outage?",
        "How would you migrate to a custom backend?",
        "How would you implement rate limiting?",
        "How would you prevent spam complaints?",
        "How would you improve AI categorization?",
        "How would you ensure accessibility?",
        "What are security risks in client-side localStorage?",
        "What is the difference between RLS and API-level authorization?",
        "Why is server-side validation still needed if frontend validates?",
        "How would you protect against XSS?",
        "How would you protect against CSRF?",
        "How would you measure system success?",
        "What metrics would management care about?",
        "How would you handle anonymous complaints ethically?",
        "How would you add an appeal process for rejected complaints?",
        "How would you design notification read/unread syncing?",
    ]
    p(doc, "Use this answer pattern: define the concept, connect it to your project, then mention one improvement or limitation.")
    h(doc, "Easy Questions", 2)
    nums(doc, easy)
    h(doc, "Medium Questions", 2)
    nums(doc, medium)
    h(doc, "Difficult Questions", 2)
    nums(doc, hard)

    h(doc, "Part 18 - Model Answers You Should Memorize", 1)
    table(doc, ["Topic", "Answer"], [
        ("Project in one line", "A role-based complaint management system for BZU where students submit complaints and staff manage review, assignment, resolution, and reporting."),
        ("Why needed", "It replaces manual and untracked complaint handling with transparent digital workflow."),
        ("Main innovation", "Role-based routing, complaint tracking, notifications, activity logs, and analytics in one system."),
        ("Database choice", "PostgreSQL because the system has structured relational data and reporting needs."),
        ("Security", "Supabase Auth, server-side token verification, RBAC, RLS, private storage, and environment variables."),
        ("Limitation", "Current AI classification is keyword-based and can be improved with trained NLP model."),
        ("Future work", "Mobile app, realtime notifications, SLA escalation, advanced AI, exportable reports, and better scalability."),
    ])

    h(doc, "Final Revision Checklist", 1)
    bullets(doc, [
        "Practice the 60-second introduction.",
        "Draw architecture diagram on paper at least 3 times.",
        "Explain all roles and complaint lifecycle without looking.",
        "Memorize database tables and why each exists.",
        "Prepare demo accounts for Student, HOD/DSA, Faculty, Supervisor.",
        "Prepare 2-3 sample complaints: Academic, Facilities, Urgent.",
        "Know how unauthorized access is blocked.",
        "Be honest about limitations and confident about future improvements.",
        "During viva, answer in this format: concept -> project example -> benefit.",
    ])

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    h(doc, "Last-Minute Speaking Template", 1)
    p(doc, "Sir/Ma'am, this project is designed to solve the university complaint handling problem. The main users are students and university staff. Students submit complaints, and staff roles process them according to responsibility. I used Next.js because it provides both frontend and backend API routes in one framework, and Supabase because it provides PostgreSQL database, authentication, storage, and security policies. The project focuses on transparency, accountability, and management reporting.")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
